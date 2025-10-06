from django.shortcuts import render, redirect, get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .traitement_cv import extraire_infos_depuis_cv
from django.utils import timezone
import stripe
from datetime import timedelta
from django.conf import settings
from django.http import HttpResponse
from .models import Candidat, Abonnement
from docx import Document
from io import BytesIO
import os
from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx.shared import Mm
from .forms import ContactMessageForm
from django.template.loader import render_to_string
from weasyprint import HTML
import re
import json
from collections import defaultdict
from .forms import CustomUserCreationForm
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

stripe.api_key = settings.STRIPE_SECRET_KEY

def contact_view(request):
    if request.method == "POST":
        form = ContactMessageForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Message envoyé avec succès !")
            return redirect("")  # Redirige vers la même page après envoi
    else:
        form = ContactMessageForm()

    return render(request, "dashboard/index.html", {"form": form})

def appliquer_theme(document, theme, nom, prenom, titre_poste, description, infos):
    """
    Applique un style au document Word selon le thème choisi.
    """
    if theme == "theme1":
        # Thème 1 : Minimaliste
        titre = document.add_heading(f"{prenom} {nom}", level=1)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titre.runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)  # Bleu foncé

        if titre_poste:
            p = document.add_paragraph(titre_poste)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.italic = True

        document.add_paragraph(description or "")

    elif theme == "theme2":
        # Thème 2 : Moderne avec séparateurs
        titre = document.add_heading(f"{prenom.upper()} {nom.upper()}", level=1)
        run = titre.runs[0]
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge vif

        if titre_poste:
            document.add_paragraph(f"💼 {titre_poste}")

        document.add_paragraph("──────────────────────────")

        if description:
            document.add_paragraph(f"📝 {description}")

    elif theme == "theme3":
        # Thème 3 : Élégant noir et gris
        titre = document.add_heading(f"{prenom} {nom}", level=1)
        run = titre.runs[0]
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(50, 50, 50)  # Gris foncé

        if titre_poste:
            p = document.add_paragraph(titre_poste)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        if description:
            p = document.add_paragraph(description)
            p.runs[0].font.italic = True

    else:
        # Thème par défaut
        document.add_heading(f"{prenom} {nom}", level=1)
        if titre_poste:
            document.add_paragraph(titre_poste)
        if description:
            document.add_paragraph(description)

    # Infos supplémentaires (ex. coordonnées)
    if infos:
        document.add_paragraph(f"📞 {infos}")

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from io import BytesIO

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Frame
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


from django.templatetags.static import static
from reportlab.lib.utils import ImageReader
from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from django.contrib.auth.decorators import login_required

from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML

@login_required
def telecharger_cv(request):
    if request.method != "POST":
        return HttpResponse("Méthode non autorisée", status=405)

    user = request.user

    # 🔹 Récupérer les valeurs simples
    nom = request.POST.get("nom", "").strip()
    prenom = request.POST.get("prenom", "").strip()
    titre_poste = request.POST.get("posttitle", "").strip()
    informations = request.POST.get("Informations", "").strip()
    certificats = request.POST.get("certificats", "").strip()
    experience_raw = request.POST.get("experience", "").strip()
    formation_raw = request.POST.get("formation", "").strip()  # <- nouveau champ
    langues = request.POST.get("Langues", "").strip()
    description = request.POST.get("Description", "").strip()
    modele_option = request.POST.get("modele_option", "choisir")
    rajouter = request.POST.get("rajouter", "").strip()
    theme = request.POST.get("theme", "theme1")
    infos_entreprise = request.POST.get("infos_entreprise", "non")

    # ✅ Style personnalisé
    font_family = request.POST.get("font_family", "Helvetica")
    font_size = int(request.POST.get("font_size", 13))
    color_titles = request.POST.get("color_titles", "#000000")
    color_details = request.POST.get("color_details", "#333333")

    # 🔹 Enregistrer le candidat
    candidat = Candidat.objects.create(
        user=user,
        nom=nom,
        prenom=prenom,
        titre_poste=titre_poste,
        description_profil=description,
        langues=langues,
        experience=experience_raw,
        certificats=certificats,
        competences=request.POST.get("competences", "").strip(),
        informations_contact=informations,
        formation=formation_raw,  # <- ajouté
        rajouter_champ=rajouter,
    )

    # 🔹 Charger un fond si uploadé (optionnel)
    background_path = None
    if modele_option == "upload" and request.FILES.get("modele_file"):
        uploaded_file = request.FILES["modele_file"]
        background_path = os.path.join(settings.MEDIA_ROOT, "uploads", uploaded_file.name)
        os.makedirs(os.path.dirname(background_path), exist_ok=True)

        with open(background_path, "wb+") as dest:
            for chunk in uploaded_file.chunks():
                dest.write(chunk)
    else:
        theme_img_path = os.path.join(
            settings.BASE_DIR, "dashboard", "static", "themes", f"{theme}.png"
        )
        if os.path.exists(theme_img_path):
            background_path = theme_img_path
    background_url = None
    if background_path:
        if background_path.startswith(settings.MEDIA_ROOT):
            background_url = background_path.replace(settings.MEDIA_ROOT, settings.MEDIA_URL)
        else:
            background_url = os.path.relpath(background_path, settings.BASE_DIR)

    # ------------------------
    # 🔹 Parser les contacts
    # ------------------------
    contacts = []
    pattern_contact = re.compile(r'^contact\[(\d+)\]\[(\w+)\]$')
    buckets_contact = defaultdict(dict)
    for key in request.POST.keys():
        m = pattern_contact.match(key)
        if m:
            idx = int(m.group(1))
            field = m.group(2)
            buckets_contact[idx][field] = request.POST.get(key, "").strip()
    for idx in sorted(buckets_contact.keys()):
        c = buckets_contact[idx]
        if c.get("type") or c.get("valeur"):
            contacts.append({"type": c.get("type", ""), "valeur": c.get("valeur", "")})

    # ------------------------
    # 🔹 Parser les compétences
    # ------------------------
    competences = []
    pattern_comp = re.compile(r'^competences\[(\d+)\]\[(\w+)\]$')
    buckets_comp = defaultdict(dict)
    for key in request.POST.keys():
        m = pattern_comp.match(key)
        if m:
            idx = int(m.group(1))
            field = m.group(2)
            buckets_comp[idx][field] = request.POST.get(key, "").strip()
    for idx in sorted(buckets_comp.keys()):
        comp = buckets_comp[idx]
        if comp.get("titre") or comp.get("details"):
            competences.append({
                "titre": comp.get("titre", ""),
                "details": [s.strip() for s in re.split(r'[,\n;]+', comp.get("details", "")) if s.strip()]
            })

    # ------------------------
    # 🔹 Parser l’expérience (robuste)
    # ------------------------
    experiences = []

    def _normalize_item(d):
        if not isinstance(d, dict):
            return {"poste": "", "entreprise": "", "periode": "", "details": str(d).strip()}
        return {
            "poste": str(d.get("poste") or d.get("title") or d.get("job") or d.get("poste_title") or "").strip(),
            "entreprise": str(d.get("entreprise") or d.get("company") or d.get("employer") or "").strip(),
            "periode": str(d.get("periode") or d.get("period") or d.get("dates") or d.get("year") or "").strip(),
            "details": str(d.get("details") or d.get("description") or d.get("missions") or "").strip()
        }

    parsed = False
    if experience_raw:
        try:
            candidate = json.loads(experience_raw)
            if isinstance(candidate, list):
                experiences = [_normalize_item(it) for it in candidate]
                parsed = True
            elif isinstance(candidate, dict):
                ev = candidate.get("experience") or candidate.get("experiences")
                if isinstance(ev, list):
                    experiences = [_normalize_item(it) for it in ev]
                    parsed = True
        except Exception:
            parsed = False

    if not parsed:
        pattern = re.compile(r'^experience\[(\d+)\]\[(\w+)\]$')
        buckets = defaultdict(dict)
        for key in request.POST.keys():
            m = pattern.match(key)
            if m:
                idx = int(m.group(1))
                field = m.group(2)
                buckets[idx][field] = request.POST.get(key, "").strip()
        for idx in sorted(buckets.keys()):
            experiences.append(_normalize_item(buckets[idx]))

    if not parsed and experience_raw:
        blocks = re.split(r'\n{2,}|\r\n\r\n|(?=\d{2}/\d{4})', experience_raw)
        blocks = [b.strip() for b in blocks if b.strip()]
        for b in blocks:
            experiences.append({"poste": "", "entreprise": "", "periode": "", "details": b})

    # ------------------------
    # 🔹 Parser la formation (robuste)
    # ------------------------
    formations = []

    def _normalize_formation(d):
        if not isinstance(d, dict):
            return {"diplome": "", "etablissement": "", "periode": "", "details": str(d).strip()}
        return {
            "diplome": str(d.get("diplome", "") or "").strip(),
            "etablissement": str(d.get("etablissement", "") or d.get("ecole", "") or "").strip(),
            "periode": str(d.get("periode", "") or "").strip(),
            "details": str(d.get("details", "") or "").strip()
        }

    parsed_formation = False
    if formation_raw:
        try:
            candidate = json.loads(formation_raw)
            if isinstance(candidate, list):
                formations = [_normalize_formation(it) for it in candidate]
                parsed_formation = True
            elif isinstance(candidate, dict):
                ev = candidate.get("formation") or candidate.get("formations")
                if isinstance(ev, list):
                    formations = [_normalize_formation(it) for it in ev]
                    parsed_formation = True
        except Exception:
            parsed_formation = False

    if not parsed_formation:
        pattern = re.compile(r'^formations\[(\d+)\]\[(\w+)\]$')
        buckets = defaultdict(dict)
        for key in request.POST.keys():
            m = pattern.match(key)
            if m:
                idx = int(m.group(1))
                field = m.group(2)
                buckets[idx][field] = request.POST.get(key, "").strip()
        for idx in sorted(buckets.keys()):
            formations.append(_normalize_formation(buckets[idx]))

    if not parsed_formation and formation_raw:
        blocks = re.split(r'\n{2,}|\r\n\r\n', formation_raw)
        blocks = [b.strip() for b in blocks if b.strip()]
        for b in blocks:
            formations.append({"diplome": "", "etablissement": "", "periode": "", "details": b})

    # ------------------------
    # 🔹 Structurer les autres champs
    # ------------------------
    def split_to_list(s):
        if not s:
            return []
        try:
            val = json.loads(s)
            if isinstance(val, list):
                return [str(x).strip() for x in val if str(x).strip()]
        except Exception:
            pass
        return [p.strip() for p in re.split(r'[,\n;]+', s) if p.strip()]

    data = {
        "nom": nom,
        "prenom": prenom,
        "titre_poste": titre_poste,
        "contacts": contacts,
        "competences": competences,
        "certificats": split_to_list(certificats),
        "experience": experiences,
        "formations": formations,
        "langues": split_to_list(langues),
        "description": description,
        "rajouter": rajouter,
        "font_family": font_family,
        "font_size": font_size,
        "color_titles": color_titles,
        "color_details": color_details,
        "background": background_path,
        "logo": user.logo.url if hasattr(user, "logo") and user.logo else None,
        "background_url": background_url,
        "footer": (
            f"{user.company_name or ''} | {user.company_address or ''} | {user.company_phone or ''}"
            + (f" | KBIS: {user.company_kbis}" if getattr(user, "company_kbis", None) else "")
            if infos_entreprise == "oui"
            else ""
        ),
    }

    # 🔹 Rendu HTML + PDF
    # 🔹 Choisir le template à utiliser
    if modele_option == "upload":
        template_name = "cv_templates/uploaded_template.html"
    else:
        template_name = f"cv_templates/{theme}.html"

    # 🔹 Rendu HTML + PDF
    html_content = render_to_string(template_name, data)

    pdf_file = HTML(string=html_content, base_url=request.build_absolute_uri("/")).write_pdf()

    # 🔹 Stats
    user.cv_download_count += 1
    user.nombre_cv_total += 1
    user.nombre_cv_peut_telecharger -= 1
    debut_mois = user.abonnement_debut.replace(day=1)
    if candidat.date_creation.date() >= debut_mois:
        user.nombre_cv_mois += 1
    user.save()

    filename = f"{user.company_name or 'Entreprise'}_{nom or ''}_{prenom or ''}.pdf"
    response = HttpResponse(pdf_file, content_type="application/pdf")
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

@login_required
def abonnement_view(request):
    user = request.user
    today = timezone.now().date()
    abonnements = Abonnement.objects.all()

    return render(request, "dashboard/abonnement.html", {
        "abonnements": abonnements,
        "stripe_pub_key": settings.STRIPE_PUBLISHABLE_KEY, 
        "user": user, 
        "today": today
    })

import stripe
from django.conf import settings
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt

@login_required
@csrf_exempt
def create_checkout_session(request):
    if request.method == 'POST':
        try:
            prix = int(request.POST.get("prix"))
            nom_abonnement = request.POST.get("nom_abonnement")

            # Détecter le domaine selon l'environnement
            if settings.DEBUG:
                domain = "http://127.0.0.1:8000"
            else:
                domain = "https://necform.onrender.com"

            stripe.api_key = settings.STRIPE_SECRET_KEY

            # Créer la session Stripe
            session = stripe.checkout.Session.create(
                payment_method_types=['card'],
                line_items=[{
                    'price_data': {
                        'currency': 'eur',
                        'product_data': {'name': nom_abonnement},
                        'unit_amount': prix,
                    },
                    'quantity': 1,
                }],
                mode='payment',
                success_url=f"{domain}/dashboard/success/?abonnement_type={nom_abonnement}",
                cancel_url=f"{domain}/dashboard/abonnement/",
                customer_email=request.user.email
            )

            return JsonResponse({'id': session.id})

        except Exception as e:
            print("Erreur Stripe :", e)
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Méthode invalide'}, status=400)


@login_required
def success_payment(request, duree = 30):
    user = request.user
    abonnement_type = request.GET.get("abonnement_type", "Basique")
    user.abonnement_type = abonnement_type
    user.abonnement_debut = timezone.now()
    abonnement_new = Abonnement.objects.get(type=user.abonnement_type)
    nombre_cv_limite = abonnement_new.nombre_cv
    nb_restant = max(user.nombre_cv_peut_telecharger, 0)  # ce qui reste avant paiement
    user.nombre_cv_peut_telecharger = nb_restant + nombre_cv_limite
    today = timezone.now().date()
    if not user.abonnement_fin or user.abonnement_fin < timezone.now().date():
        user.abonnement_fin = timezone.now().date() + timedelta(days=duree)
    else:
        user.abonnement_fin = user.abonnement_fin + timedelta(days=duree)

    user.save()
    return render(request, "dashboard/success.html", {"user": user, "today": today, })

@login_required
def cancel_payment(request):
    user = request.user
    today = timezone.now().date()
    return render(request, "dashboard/cancel.html", {"user": user, "today": today})

def index(request):
    user = request.user
    today = timezone.now().date()
    abonnement_valide = False
    if user.abonnement_fin and user.abonnement_fin >= timezone.now().date():
        abonnement_valide = True
    return render(request, 'dashboard/importcv.html', {"user": user, "abonnement_valide": abonnement_valide, "today": today})

def logout_view(request):
    logout(request)
    return redirect('login')

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            return redirect('dashboard')  
        else:
            messages.error(request, 'Nom d’utilisateur ou mot de passe incorrect.')
    
    return render(request, 'dashboard/login.html')  

@csrf_exempt
@login_required
def scanner_cv(request):
    if request.user.abonnement_fin < timezone.now().date():
        return JsonResponse({
            "success": False,
            "error": "Votre abonnement a expiré. Veuillez le renouveler pour générer un CV."
        })

    if request.method == "POST":
        cv_file = request.FILES.get("cv_file")
        if not cv_file:
            return JsonResponse({"success": False, "error": "Aucun fichier reçu."})

        data = extraire_infos_depuis_cv(cv_file)
        if not data:
            return JsonResponse({"success": False, "error": "Impossible d'extraire les infos."})

        # 🔹 Normaliser l'expérience
        experiences = data.get("experience", [])
        normalized_exp = []
        for exp in experiences:
            if isinstance(exp, dict):
                normalized_exp.append({
                    "poste": exp.get("poste", ""),
                    "entreprise": exp.get("entreprise", ""),
                    "lieu": exp.get("lieu", ""),
                    "periode": exp.get("periode", ""),
                    "details": exp.get("details", "")
                })
            else:
                normalized_exp.append({
                    "poste": exp,
                    "entreprise": "",
                    "lieu": "",
                    "periode": "",
                    "details": ""
                })
        data["experience"] = normalized_exp

        # 🔹 Normaliser les formations
        formations = data.get("formations", [])
        normalized_form = []
        for f in formations:
            if isinstance(f, dict):
                normalized_form.append({
                    "diplome": f.get("diplome", ""),
                    "ecole": f.get("ecole", ""),
                    "lieu": f.get("lieu", ""),
                    "periode": f.get("periode", ""),
                    "details": f.get("details", "")
                })
        data["formations"] = normalized_form

        # 🔹 Normaliser les compétences
        competences = data.get("competences", [])
        normalized_comp = []
        for comp in competences:
            if isinstance(comp, dict):
                details = comp.get("details", [])
                if isinstance(details, str):
                    details = [d.strip() for d in re.split(r"[,\n;]+", details) if d.strip()]
                normalized_comp.append({
                    "titre": comp.get("titre", ""),
                    "details": details
                })
            elif isinstance(comp, str):
                normalized_comp.append({"titre": comp, "details": []})
        data["competences"] = normalized_comp

        # 🔹 Normaliser les contacts
        contacts = data.get("contact", [])
        normalized_contact = []
        if isinstance(contacts, list):
            for c in contacts:
                if isinstance(c, dict):
                    normalized_contact.append({
                        "type": c.get("type", "autre"),
                        "valeur": c.get("valeur", "")
                    })
                elif isinstance(c, str):
                    # heuristique
                    if "@" in c:
                        normalized_contact.append({"type": "email", "valeur": c})
                    elif any(ch.isdigit() for ch in c):
                        normalized_contact.append({"type": "téléphone", "valeur": c})
                    else:
                        normalized_contact.append({"type": "autre", "valeur": c})
        elif isinstance(contacts, str):
            normalized_contact.append({"type": "autre", "valeur": contacts})
        data["contact"] = normalized_contact

        return JsonResponse({"success": True, **data})

    return JsonResponse({"success": False, "error": "Méthode invalide."})

@login_required
def candidats_list(request):
    user = request.user
    today = timezone.now().date()
    candidats = request.user.candidats.all()
    return render(request, "dashboard/candidats_list.html", {"candidats": candidats, "user": user, "today": today})

@login_required
def dashboard_view(request):
    user = request.user
    today = timezone.now().date()

    # --- Récupération de l'abonnement ---
    abonnement = None
    nombre_cv_limite = 0
    if user.abonnement_type:
        try:
            abonnement = Abonnement.objects.get(type=user.abonnement_type)
            nombre_cv_limite = abonnement.nombre_cv
        except Abonnement.DoesNotExist:
            abonnement = None
            nombre_cv_limite = 0

    # --- Calcul du nombre de CV restant ---
    cv_telecharges = user.cv_download_count
    cv_restants = max(nombre_cv_limite - cv_telecharges, 0)

    # --- Vérification si l'utilisateur peut télécharger ---
    can_download = cv_restants > 0

    # --- Mise à jour des infos entreprise ---
    if request.method == "POST":
        user.company_name = request.POST.get("company_name", user.company_name)
        user.company_address = request.POST.get("company_address", user.company_address)
        user.company_phone = request.POST.get("company_phone", user.company_phone)
        user.company_kbis = request.POST.get("company_kbis", user.company_kbis)

        if request.FILES.get("logo"):
            user.logo = request.FILES["logo"]

        user.save()
        return redirect("dashboard")  # Recharge la page après mise à jour

    return render(request, "dashboard/dashboard.html", {
        "user": user,
        "today": today,
        "abonnement": abonnement,
        "cv_telecharges": cv_telecharges,
        "cv_restants": cv_restants,
        "can_download": can_download,
    })

def signup_view(request):
    if request.method == "POST":
        form = CustomUserCreationForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, "Your account has been created successfully! Please log in.")
            return redirect("login")  
    else:
        form = CustomUserCreationForm()
    return render(request, "signup.html", {"form": form})
